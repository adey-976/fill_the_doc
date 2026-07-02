"""
Replace placeholder tags in a .docx file with content from .md files.

Usage:
    python3 replace_tags_in_docx.py XYZ_PostProcessing_ODD.docx

Output:
    <name>_ORR.docx  (the mapped tags replaced by the corresponding .md content)

Supported markdown: headings (#), paragraphs, bullet/ordered lists, tables
(with **bold** header cells and <br> line breaks), horizontal rules, and inline
**bold** / *italic* / `code`. Each markdown element must be on its own line.
"""

import sys
import re
import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TAG_FILE_MAPPING = {
    "<review_summary_and_limitation_table>": "summary_limitation.md",
    "<section_1_reviewer_assessment>": "sec_1.md",
    "<section_2_reviewer_assessment>": "sec_2.md",
    "<section_3_reviewer_assessment>": "sec_3.md",
}


def add_runs(para, text):
    """Add runs to a paragraph, honoring **bold**, *italic* and `code`."""
    for chunk in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            para.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            para.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            para.add_run(chunk[1:-1]).font.name = "Courier New"
        else:
            para.add_run(chunk)


def add_table(doc, rows):
    """Build a Table Grid table; first row and fully-bold cells are bold."""
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            bold = text.startswith("**") and text.endswith("**")
            for k, line in enumerate(text.replace("<br>", "\n").split("\n")):
                para = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                add_runs(para, line)
                for run in para.runs:
                    run.font.size = Pt(9)
                    if i == 0 or bold:
                        run.bold = True
    return table


def parse_table(lines):
    """Turn markdown table lines into rows of cells, dropping the separator row."""
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in lines]
    if len(rows) >= 2 and all(re.fullmatch(r"[:\-\s]*", c) for c in rows[1]):
        rows = [rows[0]] + rows[2:]
    return rows


def hr_paragraph(doc):
    """A blank paragraph with a bottom border (horizontal rule)."""
    para = doc.add_paragraph()
    bottom = OxmlElement("w:bottom")
    for attr, val in (("val", "single"), ("sz", "6"), ("space", "1"), ("color", "auto")):
        bottom.set(qn(f"w:{attr}"), val)
    pbdr = OxmlElement("w:pBdr")
    pbdr.append(bottom)
    para._p.get_or_add_pPr().append(pbdr)
    return para


def md_to_elements(md_text, doc):
    """Parse markdown into docx body elements built in `doc` (a scratch document)."""
    elements = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        # Table: consecutive lines starting with "|"
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            elements.append(add_table(doc, parse_table(block))._tbl)
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 9)
            para = doc.add_paragraph(style=f"Heading {level}")
            add_runs(para, stripped.lstrip("#").strip())
            elements.append(para._p)

        elif stripped in ("---", "***", "___"):
            elements.append(hr_paragraph(doc)._p)

        elif re.match(r"[*+-]\s+", stripped) or re.match(r"\d+[.)]\s+", stripped):
            ordered = bool(re.match(r"\d+[.)]\s+", stripped))
            level = (len(line) - len(line.lstrip())) // 4
            text = re.sub(r"^(\d+[.)]|[*+-])\s+", "", stripped)
            para = doc.add_paragraph(style="List Paragraph")
            para.paragraph_format.left_indent = Pt(18 * (level + 1))
            para.add_run(f"{stripped.split()[0]} " if ordered else "•  ")
            add_runs(para, text)
            elements.append(para._p)

        else:
            para = doc.add_paragraph()
            add_runs(para, stripped)
            elements.append(para._p)

        i += 1

    return elements


def find_and_replace_tags(input_path):
    input_path = Path(input_path)
    if not input_path.exists():
        sys.exit(f"Error: Input file '{input_path}' not found.")

    output_path = input_path.with_name(f"{input_path.stem}_ORR{input_path.suffix}")
    doc = Document(str(input_path))

    targets = [p for p in doc.paragraphs if p.text.strip() in TAG_FILE_MAPPING]
    if not targets:
        print("No tags found in the document.")
        return

    for para in targets:
        md_path = input_path.parent / TAG_FILE_MAPPING[para.text.strip()]
        if not md_path.exists():
            print(f"Warning: '{md_path.name}' not found; leaving tag in place.")
            continue

        elements = md_to_elements(md_path.read_text(encoding="utf-8"), Document())
        anchor = para._element
        parent = anchor.getparent()
        idx = list(parent).index(anchor)
        for elem in elements:
            parent.insert(idx, copy.deepcopy(elem))
            idx += 1
        parent.remove(anchor)
        print(f"Replaced '{para.text.strip()}' with '{md_path.name}'.")

    doc.save(str(output_path))
    print(f"\nOutput saved to: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("Usage: python3 replace_tags_in_docx.py <input_file.docx>")
    find_and_replace_tags(sys.argv[1])
