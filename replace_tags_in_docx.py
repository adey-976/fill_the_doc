"""
Replace placeholder tags in a .docx file with content from .md files.

Usage:
    python replace_tags_in_docx.py XYZ_PostProcessing_ODD.docx

Output:
    Creates a copy with "_ORR" appended to the filename (before extension),
    with all mapped tags replaced by the corresponding .md file content.
"""

import sys
import re
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# TAG-TO-FILE MAPPING
# =============================================================================
# Edit this dictionary to add/update/remove tag-to-file mappings.
# Keys are the tag strings (including angle brackets) as they appear in the doc.
# Values are the .md filenames (looked up in the same directory as the input).
# =============================================================================

TAG_FILE_MAPPING = {
    "<review_summary_and_limitation_table>": "summary_limitation.md",
    "<section_1_reviewer_assessment>": "sec_1.md",
    "<section_2_reviewer_assessment>": "sec_2.md",
    "<section_3_reviewer_assessment>": "sec_3.md",
}


# =============================================================================
# FLAT MARKDOWN PARSER (handles single-line markdown with no newlines)
# =============================================================================

def parse_flat_markdown(text):
    """Parse a flat (single-line, no newlines) markdown string into structured blocks.

    Strategy (TABLE-FIRST):
    1. Find all tables by locating separator rows (| :--- | :--- |).
    2. Extract tables with their positions in the text.
    3. Process remaining text between tables for other elements.
    """
    text = text.strip()
    blocks = []

    # Step 1: Find separator rows
    separator_pattern = re.compile(r'\|(?:\s*:?-{2,}:?\s*\|)+')
    separators = list(separator_pattern.finditer(text))

    # Step 2: Extract tables
    tables = []
    consumed_sep_indices = set()

    for sep_idx, sep_match in enumerate(separators):
        if sep_idx in consumed_sep_indices:
            continue

        sep_start = sep_match.start()
        sep_end = sep_match.end()
        sep_text = sep_match.group()

        sep_cells = [c.strip() for c in sep_text.strip('|').split('|')]
        num_cols = len(sep_cells)
        if num_cols < 1:
            continue

        # Check for dual-table (side-by-side) pattern
        dual_table = False
        if sep_idx + 1 < len(separators):
            next_sep = separators[sep_idx + 1]
            gap_text = text[sep_end:next_sep.start()]
            if gap_text.strip() == '' and len(gap_text) <= 3:
                dual_table = True
                consumed_sep_indices.add(sep_idx + 1)
                next_sep_cells = [c.strip() for c in next_sep.group().strip('|').split('|')]
                combined_cols = num_cols + len(next_sep_cells)
                next_sep_end = next_sep.end()

                needed_pipes = combined_cols + 2
                pos = sep_start - 1
                while pos > 0 and text[pos] == ' ':
                    pos -= 1
                pipe_count = 0
                header_start = None
                while pos >= 0:
                    if text[pos] == '|':
                        pipe_count += 1
                        if pipe_count == needed_pipes:
                            header_start = pos
                            break
                    pos -= 1

                if header_start is not None:
                    header_text = text[header_start:sep_start].rstrip()
                    h_parts = header_text.split('|')
                    h_cells = [c.strip() for c in h_parts if c.strip() != '' or len(h_parts) > combined_cols + 1]
                    h_cells = [c.strip() for c in h_parts[1:-1]]

                    data_rows = []
                    pos = next_sep_end
                    while pos < len(text):
                        if text[pos] == ' ':
                            pos += 1
                            continue
                        if text[pos] == '|':
                            pipe_positions = []
                            scan = pos
                            while scan < len(text):
                                if text[scan] == '|':
                                    pipe_positions.append(scan)
                                    if len(pipe_positions) == needed_pipes:
                                        break
                                scan += 1
                            if len(pipe_positions) == needed_pipes:
                                row_text = text[pos:pipe_positions[-1]+1]
                                r_parts = row_text.split('|')
                                r_cells = [c.strip() for c in r_parts[1:-1]]
                                if all(re.match(r'^:?-+:?$', c) for c in r_cells if c):
                                    break
                                data_rows.append(r_cells)
                                pos = pipe_positions[-1] + 1
                            else:
                                break
                        else:
                            break
                    tables.append((header_start, pos, h_cells, data_rows))
                    continue

        if dual_table:
            continue

        # Standard table: find header row backwards from separator
        before_sep = sep_start
        if before_sep > 0 and text[before_sep - 1] == ' ':
            before_sep -= 1

        pos = before_sep - 1
        pipe_count = 0
        header_start = None
        hit_boundary = False
        while pos >= 0:
            if text[pos] == '|':
                pipe_count += 1
                if pipe_count == num_cols + 1:
                    header_start = pos
                    break
            if text[pos] == '*' and pos >= 2 and text[pos-2:pos+1] == '***':
                hit_boundary = True
                break
            if text[pos] == '-' and pos >= 2 and text[pos-2:pos+1] == '---':
                hit_boundary = True
                break
            pos -= 1

        if header_start is None and not hit_boundary:
            continue

        if hit_boundary or header_start is None:
            h_cells = ['' for _ in range(num_cols)]
            search_start = pos + 1 if hit_boundary else 0
            first_pipe = text.find('|', search_start, sep_start)
            header_start = first_pipe if first_pipe >= 0 else sep_start
        else:
            header_text = text[header_start:before_sep]
            h_parts = header_text.split('|')
            h_cells = [c.strip() for c in h_parts[1:-1]] if len(h_parts) > 2 else [c.strip() for c in h_parts if c.strip()]
            if len(h_cells) != num_cols:
                h_cells = [c.strip() for c in h_parts if c.strip()][:num_cols]

        # Find data rows forward from separator
        data_rows = []
        pos = sep_end

        if num_cols == 2:
            data_rows, pos = _parse_2col_rows(text, pos)
        else:
            while pos < len(text):
                if text[pos] == ' ':
                    pos += 1
                    continue
                if text[pos] == '|':
                    pipe_positions = []
                    scan = pos
                    while scan < len(text):
                        if text[scan] == '|':
                            pipe_positions.append(scan)
                            if len(pipe_positions) == num_cols + 1:
                                break
                        scan += 1
                    if len(pipe_positions) == num_cols + 1:
                        row_text = text[pos:pipe_positions[-1]+1]
                        r_parts = row_text.split('|')
                        r_cells = [c.strip() for c in r_parts[1:-1]]
                        non_empty = [c for c in r_cells if c.strip()]
                        if non_empty and all(re.match(r'^:?-+:?$', c.strip()) for c in non_empty):
                            break
                        data_rows.append(r_cells[:num_cols])
                        pos = pipe_positions[-1] + 1
                    else:
                        break
                else:
                    break

        tables.append((header_start, pos, h_cells, data_rows))

    # Remove overlapping tables
    tables.sort(key=lambda t: t[0])
    filtered = []
    for t in tables:
        if filtered and t[0] < filtered[-1][1]:
            if len(t[3]) > len(filtered[-1][3]):
                filtered[-1] = t
        else:
            filtered.append(t)
    tables = filtered

    # Step 3: Process segments between tables
    segments = []
    prev_end = 0
    for t in tables:
        if t[0] > prev_end:
            segments.append(('text', text[prev_end:t[0]]))
        segments.append(('table', t))
        prev_end = t[1]
    if prev_end < len(text):
        segments.append(('text', text[prev_end:]))

    for seg_type, seg_data in segments:
        if seg_type == 'table':
            _, _, header_cells, data_rows = seg_data
            blocks.append({'type': 'table', 'header': header_cells, 'rows': data_rows})
        else:
            seg_text = seg_data.strip()
            if seg_text:
                _parse_text_segment(seg_text, blocks)

    return blocks


def _parse_2col_rows(text, pos):
    """Parse data rows for 2-column tables (handles empty cells)."""
    data_rows = []
    row_pattern = re.compile(r'\|\s([^|]*)\|\s([^|]*)\|')

    orig_pos = pos
    while pos < len(text) and text[pos] == ' ':
        pos += 1

    if pos < len(text) and text[pos] != '|':
        backup = orig_pos - 1
        if backup >= 0 and text[backup] == '|':
            pos = backup
        else:
            return data_rows, orig_pos

    while pos < len(text):
        if text[pos] == ' ':
            pos += 1
            continue
        if text[pos] != '|':
            break

        m = row_pattern.match(text, pos)
        if m:
            cell1 = m.group(1).strip()
            cell2 = m.group(2).strip()

            if cell1 and cell2 and re.match(r'^:?-+:?$', cell1) and re.match(r'^:?-+:?$', cell2):
                break
            if '***' in cell1 or '***' in cell2:
                break

            if cell1 == '' and cell2.startswith('**'):
                if pos + 2 < len(text) and text[pos+2] == '|':
                    m2 = row_pattern.match(text, pos + 2)
                    if m2:
                        cell1 = m2.group(1).strip()
                        cell2 = m2.group(2).strip()
                        if '***' in cell1 or '***' in cell2:
                            break
                        data_rows.append([cell1, cell2])
                        pos = m2.end()
                        continue
                data_rows.append([cell1, cell2])
                pos = m.end()
            elif cell1 == '' and cell2 == '':
                pos = m.end()
            else:
                data_rows.append([cell1, cell2])
                pos = m.end()
        else:
            break

    return data_rows, pos


def _parse_text_segment(text, blocks):
    """Parse a non-table text segment into headings, paragraphs, bullets, etc."""
    markers = []

    # Find headings
    for m in re.finditer(r'(?:(?<=^)|(?<= ))(#{2,6}) ', text):
        markers.append((m.start(), 'heading', m))

    # Find horizontal rules (*** or --- standalone, not inside words)
    for m in re.finditer(r'(?:(?<=^)|(?<= ))(\*{3}|-{3})(?= |$)', text):
        if m.start() > 0 and text[m.start()-1] == '|':
            continue
        markers.append((m.start(), 'hr', m))

    # Find bullet items (* followed by space, not **)
    for m in re.finditer(r'(?:(?<=^)|(?<= ))\* (?!\*)', text):
        # Exclude "- " inside words like "Summary - Review"
        # A real bullet: preceded by start, sentence-end (. : ), or another bullet
        pos = m.start()
        if pos > 0:
            before_char = text[pos-1]
            if before_char != ' ':
                continue
            # Look further back - if preceded by a word char (not punctuation), skip
            lookback = text[max(0, pos-3):pos].rstrip()
            if lookback and lookback[-1].isalpha():
                continue
        markers.append((m.start(), 'bullet', m))

    # Find ordered list items
    for m in re.finditer(r'(?:(?<=^)|(?<= ))(\d+)\. ', text):
        pos = m.start()
        if pos > 0:
            before = text[max(0, pos-3):pos].rstrip()
            if before and before[-1].isalpha():
                continue
        markers.append((m.start(), 'ordered', m))

    markers.sort(key=lambda x: x[0])

    if not markers:
        if text.strip() and text.strip() not in ('|', '| |', '||'):
            blocks.append({'type': 'paragraph', 'text': text.strip()})
        return

    # Text before first marker
    if markers[0][0] > 0:
        para_text = text[:markers[0][0]].strip()
        if para_text and para_text not in ('|', '| |'):
            blocks.append({'type': 'paragraph', 'text': para_text})

    i = 0
    while i < len(markers):
        pos, mtype, match = markers[i]
        next_pos = markers[i+1][0] if i+1 < len(markers) else len(text)

        if mtype == 'heading':
            level = len(match.group(1))
            content = text[pos + level + 1:next_pos].strip()
            blocks.append({'type': 'heading', 'level': level, 'text': content})
            i += 1

        elif mtype == 'hr':
            blocks.append({'type': 'hr'})
            remaining = text[pos + 3:next_pos].strip()
            if remaining and remaining not in ('|', '| |'):
                blocks.append({'type': 'paragraph', 'text': remaining})
            i += 1

        elif mtype == 'bullet':
            items = []
            while i < len(markers) and markers[i][1] == 'bullet':
                bpos = markers[i][0]
                bnext = markers[i+1][0] if i+1 < len(markers) else len(text)
                item_text = text[bpos+2:bnext].strip()
                items.append({'text': item_text, 'level': 0})
                i += 1
            blocks.append({'type': 'bullet_list', 'items': items})

        elif mtype == 'ordered':
            items = []
            while i < len(markers) and markers[i][1] == 'ordered':
                opos = markers[i][0]
                onext = markers[i+1][0] if i+1 < len(markers) else len(text)
                marker_match = re.match(r'\d+\.\s', text[opos:])
                mlen = marker_match.end() if marker_match else 3
                item_text = text[opos + mlen:onext].strip()
                items.append({'text': item_text, 'level': 0})
                i += 1
            blocks.append({'type': 'ordered_list', 'items': items})

        else:
            i += 1


# =============================================================================
# DOCX RENDERING
# =============================================================================

def add_formatted_runs(paragraph, text):
    """Add runs to a paragraph with bold/italic/code markdown formatting."""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        if match.group(2):
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):
            run = paragraph.add_run(match.group(5))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        last_end = match.end()
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def render_table(doc, header, rows):
    """Create a Word table from parsed header and rows."""
    num_cols = max(len(header), max((len(r) for r in rows), default=0)) if rows else len(header)
    all_rows = [header] + rows
    table = doc.add_table(rows=len(all_rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(all_rows):
        for j in range(num_cols):
            cell = table.rows[i].cells[j]
            cell_text = row_data[j] if j < len(row_data) else ""
            clean_text = cell_text.strip().replace("<br>", "\n")

            cell.paragraphs[0].clear()
            for k, line in enumerate(clean_text.split("\n")):
                para = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.space_before = Pt(2)
                add_formatted_runs(para, line)
                if i == 0:
                    for run in para.runs:
                        run.bold = True
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9)

    return table


def create_list_paragraph(text, bullet=True, bullet_num_id="2", ordered_num_id="1"):
    """Create a list paragraph element with proper numbering."""
    para_elem = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "ListParagraph")
    pPr.append(pStyle)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numPr.append(ilvl)
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), bullet_num_id if bullet else ordered_num_id)
    numPr.append(numId_elem)
    pPr.append(numPr)
    para_elem.append(pPr)

    # Add formatted text runs
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            _add_run(para_elem, text[last_end:match.start()])
        if match.group(2):
            _add_run(para_elem, match.group(2), bold=True, italic=True)
        elif match.group(3):
            _add_run(para_elem, match.group(3), bold=True)
        elif match.group(4):
            _add_run(para_elem, match.group(4), italic=True)
        elif match.group(5):
            _add_run(para_elem, match.group(5), code=True)
        last_end = match.end()
    if last_end < len(text):
        _add_run(para_elem, text[last_end:])

    return para_elem


def _add_run(para_elem, text, bold=False, italic=False, code=False):
    """Add a formatted run to a paragraph element."""
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if code:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)
    run_elem.append(rPr)
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    run_elem.append(t_elem)
    para_elem.append(run_elem)


def detect_numbering_ids(doc):
    """Detect bullet/ordered list numId values from the document's numbering definitions."""
    bullet_num_id = "2"
    ordered_num_id = "1"
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part:
            numbering = numbering_part.numbering_definitions._numbering
            abs_format = {}
            for an in numbering.findall(qn("w:abstractNum")):
                an_id = an.get(qn("w:abstractNumId"))
                lvl0 = an.find(qn("w:lvl"))
                if lvl0 is not None:
                    fmt = lvl0.find(qn("w:numFmt"))
                    if fmt is not None:
                        abs_format[an_id] = fmt.get(qn("w:val"))
            for n in numbering.findall(qn("w:num")):
                num_id = n.get(qn("w:numId"))
                abs_id_elem = n.find(qn("w:abstractNumId"))
                if abs_id_elem is not None:
                    fmt = abs_format.get(abs_id_elem.get(qn("w:val")), "")
                    if fmt == "bullet":
                        bullet_num_id = num_id
                    elif fmt in ("decimal", "lowerLetter", "lowerRoman"):
                        ordered_num_id = num_id
    except Exception:
        pass
    return bullet_num_id, ordered_num_id


def insert_blocks_at_paragraph(doc, para_index, blocks):
    """Replace the paragraph at para_index with rendered markdown blocks."""
    bullet_num_id, ordered_num_id = detect_numbering_ids(doc)

    target_para = doc.paragraphs[para_index]
    target_element = target_para._element
    parent = target_element.getparent()

    new_elements = []

    for block in blocks:
        if block["type"] == "heading":
            # Normal paragraph with bold text (no Heading styles)
            new_para = OxmlElement("w:p")
            _add_run(new_para, block["text"], bold=True)
            new_elements.append(new_para)

        elif block["type"] == "paragraph":
            temp_doc = Document()
            temp_para = temp_doc.add_paragraph()
            add_formatted_runs(temp_para, block["text"])
            new_elements.append(copy.deepcopy(temp_para._element))

        elif block["type"] == "bullet_list":
            for item in block["items"]:
                elem = create_list_paragraph(item["text"], bullet=True,
                                             bullet_num_id=bullet_num_id, ordered_num_id=ordered_num_id)
                new_elements.append(elem)

        elif block["type"] == "ordered_list":
            for item in block["items"]:
                elem = create_list_paragraph(item["text"], bullet=False,
                                             bullet_num_id=bullet_num_id, ordered_num_id=ordered_num_id)
                new_elements.append(elem)

        elif block["type"] == "table":
            temp_doc = Document()
            tbl = render_table(temp_doc, block["header"], block["rows"])
            if tbl:
                new_elements.append(copy.deepcopy(tbl._tbl))

        elif block["type"] == "hr":
            temp_doc = Document()
            temp_para = temp_doc.add_paragraph()
            pPr = temp_para._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)
            new_elements.append(copy.deepcopy(temp_para._element))

    for elem in new_elements:
        parent.insert(list(parent).index(target_element), elem)
    parent.remove(target_element)


# =============================================================================
# MAIN
# =============================================================================

def find_and_replace_tags(input_path):
    """Find tags in the docx and replace them with rendered .md content."""
    input_path = Path(input_path)
    if not input_path.exists():
        print(f"Error: Input file '{input_path}' not found.")
        sys.exit(1)

    output_path = input_path.parent / f"{input_path.stem}_ORR{input_path.suffix}"
    doc = Document(str(input_path))

    replacements = []
    for i, para in enumerate(doc.paragraphs):
        full_text = para.text.strip()
        if full_text in TAG_FILE_MAPPING:
            md_filename = TAG_FILE_MAPPING[full_text]
            md_path = input_path.parent / md_filename
            if not md_path.exists():
                print(f"Warning: Markdown file '{md_path}' not found for tag '{full_text}'. Skipping.")
                continue
            replacements.append((i, full_text, md_path))

    if not replacements:
        print("No tags found in the document.")
        return

    print(f"Found {len(replacements)} tag(s) to replace:")
    for idx, tag, md_path in replacements:
        print(f"  Paragraph {idx}: '{tag}' -> '{md_path.name}'")

    for idx, tag, md_path in reversed(replacements):
        md_content = md_path.read_text(encoding="utf-8")
        blocks = parse_flat_markdown(md_content)
        insert_blocks_at_paragraph(doc, idx, blocks)
        print(f"  Replaced '{tag}' with content from '{md_path.name}' ({len(blocks)} blocks)")

    doc.save(str(output_path))
    print(f"\nOutput saved to: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python replace_tags_in_docx.py <input_file.docx>")
        print("  The script looks for .md files in the same directory as the input file.")
        sys.exit(1)

    find_and_replace_tags(sys.argv[1])
